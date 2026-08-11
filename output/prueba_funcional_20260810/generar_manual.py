
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

ROOT=Path(r"C:\Desarrollo\Desarrollo\FullStack\EGestion360")
CAP=ROOT/"output"/"prueba_funcional_20260810"/"capturas"
OUT=ROOT/"output"/"prueba_funcional_20260810"
OUT.mkdir(parents=True,exist_ok=True)
DOCX=OUT/"Manual_visual_prueba_funcional_completa.docx"

doc=Document()
sec=doc.sections[0]
sec.top_margin=Inches(.55); sec.bottom_margin=Inches(.55)
sec.left_margin=Inches(.65); sec.right_margin=Inches(.65)
sec.header_distance=Inches(.25); sec.footer_distance=Inches(.25)

styles=doc.styles
styles["Normal"].font.name="Aptos"; styles["Normal"].font.size=Pt(9)
styles["Title"].font.name="Aptos Display"; styles["Title"].font.size=Pt(30); styles["Title"].font.bold=True; styles["Title"].font.color.rgb=RGBColor(21,54,93)
for nm,size,color in [("Heading 1",20,(21,54,93)),("Heading 2",14,(28,92,140)),("Heading 3",11,(45,119,156))]:
    styles[nm].font.name="Aptos Display"; styles[nm].font.size=Pt(size); styles[nm].font.bold=True; styles[nm].font.color.rgb=RGBColor(*color)
if "Paso" not in styles:
    st=styles.add_style("Paso",WD_STYLE_TYPE.PARAGRAPH)
    st.font.name="Aptos Display"; st.font.size=Pt(12); st.font.bold=True; st.font.color.rgb=RGBColor(21,54,93)
    st.paragraph_format.space_before=Pt(7); st.paragraph_format.space_after=Pt(3)

def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement("w:shd"); shd.set(qn("w:fill"),fill); tcPr.append(shd)
def border_cell(cell,color="D9E2F3"):
    tcPr=cell._tc.get_or_add_tcPr(); borders=tcPr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders=OxmlElement("w:tcBorders"); tcPr.append(borders)
    for edge in ("top","left","bottom","right"):
        tag=OxmlElement("w:"+edge); tag.set(qn("w:val"),"single"); tag.set(qn("w:sz"),"4"); tag.set(qn("w:color"),color); borders.append(tag)
def table(headers,rows,widths=None):
    t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=True
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; shade(c,"15365D"); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs: r.font.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(8)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v); border_cell(cells[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(1)
                for r in p.runs:r.font.size=Pt(8)
    return t
def pill(text,fill="DDEBF7",color=(21,54,93)):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=t.cell(0,0); c.text=text; shade(c,fill); border_cell(c,fill)
    p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:r.font.bold=True;r.font.color.rgb=RGBColor(*color);r.font.size=Pt(10)
def note(title,text,fill="EAF3F8",color=(21,54,93)):
    t=doc.add_table(rows=1,cols=1); c=t.cell(0,0); shade(c,fill); border_cell(c,fill)
    p=c.paragraphs[0]; r=p.add_run(title+"  ");r.bold=True;r.font.color.rgb=RGBColor(*color)
    p.add_run(text)
def image(name,caption):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(CAP/name),width=Inches(7.0))
    p2=doc.add_paragraph(caption); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in p2.runs:r.italic=True;r.font.size=Pt(8);r.font.color.rgb=RGBColor(89,89,89)
def step(n,title,menu,action,result,img=None):
    doc.add_paragraph(f"Paso {n}. {title}",style="Paso")
    table(["Dónde ir","Qué hacer","Resultado esperado"],[[menu,action,result]])
    if img:image(img,f"Evidencia {n}. {title}")
def page():
    doc.add_page_break()

# Header/footer
header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
rh=header.add_run("PCI · Manual visual de prueba funcional");rh.bold=True;rh.font.color.rgb=RGBColor(21,54,93);rh.font.size=Pt(8)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
rf=footer.add_run("Flujo integral 2026 · Requisición a pólizas y almacén");rf.font.size=Pt(8);rf.font.color.rgb=RGBColor(100,100,100)

# Cover
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(55)
r=p.add_run("PCI");r.bold=True;r.font.size=Pt(42);r.font.color.rgb=RGBColor(21,54,93)
p=doc.add_paragraph("MANUAL VISUAL DE PRUEBA FUNCIONAL");p.alignment=WD_ALIGN_PARAGRAPH.CENTER
for r in p.runs:r.bold=True;r.font.size=Pt(15);r.font.color.rgb=RGBColor(45,119,156)
p=doc.add_paragraph("Del presupuesto autorizado a la orden de compra, factura, pólizas y entrada de almacén")
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
for r in p.runs:r.font.size=Pt(18);r.bold=True
doc.add_paragraph("")
pill("EJERCICIO 2026  ·  EMPRESA: PLATAFORMA DE COMPRAS INTEGRAL  ·  PRUEBA COMPLETA","DDEBF7")
doc.add_paragraph("")
note("Resultado:", "flujo ejecutado de punta a punta con datos reales de prueba. Se comprobaron la póliza de presupuesto autorizado, la póliza CP1 del compromiso, la póliza DEV de la factura y la entrada trazable al almacén.","E2F0D9",(46,92,42))
doc.add_paragraph("")
table(["Elemento","Dato usado"],[
["Usuario operativo","GABRIELA CORONA ESPINOSA"],
["Área","ALMACÉN TULTITLÁN"],
["Requisición","REQ-2026-000002"],
["Bien","MO00001 · Caja de cartón para archivo tamaño oficio"],
["Partida correcta","21101 · Materiales y útiles de oficina"],
["Orden","OC-2026-0001"],
["Factura","FAC-PRUEBA-0001"],
["Contrato / compromiso","AUT-1"],
])
doc.add_paragraph("")
p=doc.add_paragraph("Fecha de validación: 10 de agosto de 2026");p.alignment=WD_ALIGN_PARAGRAPH.CENTER
page()

doc.add_heading("1. Qué se comprobó",level=1)
table(["Etapa","Registro / evidencia","Estado"],[
["Presupuesto autorizado","Póliza 8 · $104,833.35","Balanceada"],
["Requisición","REQ-2026-000002 · partida 21101","Completa"],
["Cotización","3 proveedores; precios $1, $2 y $3","Comparativo"],
["Suficiencia","Solicitud 1 y autorización 1","Autorizada"],
["Compromiso","AUT-1 · póliza CP1 · $2.00","Vigente"],
["Orden de compra","OC-2026-0001 · $1.00","Por surtir"],
["Factura","FAC-PRUEBA-0001 · $2.00","Devengada"],
["Almacén","1 unidad MO00001 vinculada a la OC","Recibida"],
["Contabilidad","Pólizas 8, CP1 y DEV","Balanceadas"],
])
note("Respuesta a la duda “Honorarios”:", "No era correcto para este bien. La partida 12101 corresponde a honorarios. Para caja de cartón/material de oficina la prueba quedó en 21101: Materiales y útiles de oficina.","FFF2CC",(127,96,0))
doc.add_heading("2. Ruta mínima del usuario",level=1)
table(["Orden","Menú","Responsable habitual","Si falta algo, pedir a"],[
["1","Presupuesto > Presupuesto autorizado","Presupuesto","Área de Presupuesto"],
["2","Adquisiciones > Requisición","Área solicitante","Administrador de usuarios / Presupuesto"],
["3","Adquisiciones > Cotización","Compras","Padrón de proveedores / Almacén"],
["4","Adquisiciones > Solicitud Suficiencia","Compras / Presupuesto","Presupuesto"],
["5","Presupuesto > Estado de Contratos","Presupuesto / Compras","Administrador de permisos"],
["6","Adquisiciones > Orden de Compra","Compras","Compras"],
["7","Tesorería > Recepción de Facturas","Cuentas por pagar","Contabilidad / Presupuesto"],
["8","Almacén > Recepción de Pedidos","Almacén","Administrador de almacén"],
["9","Contabilidad > Pólizas","Contabilidad","Contabilidad"],
])
doc.add_heading("3. Configuración previa indispensable",level=1)
table(["Validación","Dónde se configura","Qué revisar"],[
["Usuario vinculado a persona","Configuración > Sistema > Usuarios","Persona activa y correcta"],
["Área del solicitante","Configuración > Sistema > Usuarios","Área asignada; persona marcada como solicitante"],
["Autorizador","Configuración > Sistema > Usuarios / área","Persona autorizadora del área"],
["Presupuesto","Presupuesto > Presupuesto autorizado","Posición 21101 con saldo"],
["Matriz contable","Contabilidad > Matriz de conversión","Cuentas Por ejercer, Comprometido y Devengado para 21101"],
["Proveedor","Configuración > Adquisiciones > Proveedores","Proveedor activo"],
["Bien","Configuración > Almacén > Bienes y servicios","MO00001 activo y unidad PIEZA"],
["Permisos","Configuración > Sistema > Roles","create/update/authorize en la ruta operativa"],
])
page()

doc.add_heading("4. Procedimiento visual paso a paso",level=1)
step(1,"Abrir requisiciones","Adquisiciones > Requisición","Confirma ejercicio 2026, empresa y sucursal; pulsa Nuevo.","Se abre el alta de requisición.","01_requisiciones_inicio.png")
step(2,"Capturar encabezado","Adquisiciones > Requisición > Nuevo","Selecciona área ALMACÉN TULTITLÁN, solicitante, tipo Bien e importe estimado.","La requisición obtiene folio y queda editable.","02_requisicion_formulario.png")
step(3,"Confirmar requisición creada","Misma pantalla","Guarda y abre REQ-2026-000002.","Folio visible y listo para clasificación.","03_requisicion_creada.png")
step(4,"Clasificar presupuesto","Editar requisición > Clasificación presupuestal","Selecciona posición presupuestal 21101. Verifica programa, fuente, TG, DI y DG.","La requisición toma la clasificación de la posición.","04_requisicion_clasificada.png")
step(5,"Agregar partida","Detalle de requisición > Partidas","Registra partida 21101 por el importe aplicable.","Partida Materiales y útiles de oficina.","05_partida_requisicion.png")
step(6,"Agregar el bien","Detalle de requisición > Bienes","Agrega MO00001, cantidad 1, unidad PIEZA.","El bien queda ligado a la partida.","06_bien_requisicion.png")
step(7,"Validar expediente","Detalle REQ-2026-000002","Comprueba clasificación, partida y bien antes de cotizar.","Requisición completa.","07_requisicion_completa.png")
page()

step(8,"Crear cotizaciones","Adquisiciones > Cotización","Genera cotizaciones para tres proveedores y registra sus condiciones.","Tres propuestas comparables.","11_tres_cotizaciones.png")
step(9,"Solicitar suficiencia","Adquisiciones > Solicitud Suficiencia","Selecciona REQ-2026-000002 y sus cotizaciones; guarda la solicitud.","Solicitud 1 por promedio de $2.00.","13_solicitud_suficiencia_creada.png")
step(10,"Autorizar suficiencia","Presupuesto > Autorización de Suficiencia","Abre la solicitud, registra autorizador y justificación; autoriza.","Suficiencia autorizada.","14_suficiencia_autorizada.png")
step(11,"Crear compromiso","Presupuesto > Estado de Contratos","Crea AUT-1 desde la suficiencia, proveedor adjudicado y detalle 21101.","Compromiso en borrador por $2.00.","15_compromiso_borrador.png")
step(12,"Autorizar compromiso","Presupuesto > Estado de Contratos > AUT-1","Autoriza el contrato/compromiso.","Estado Vigente y póliza CP1 automática.","16_compromiso_autorizado_poliza_cp1.png")
step(13,"Completar orden","Adquisiciones > Orden de Compra > OC-2026-0001","Agrega MO00001, cantidad 1, precio $1.00 y partida 21101 por $1.00.","Detalle y partida coinciden.","17_orden_compra_completa.png")
step(14,"Autorizar orden","Misma orden > Autorizar","Confirma la autorización.","La orden queda Por surtir y bloqueada.","18_orden_compra_autorizada.png")
page()

step(15,"Registrar factura","Tesorería > Cuentas por pagar > Recepción de Facturas y Comprobantes","Pulsa Nuevo, elige AUT-1; el sistema carga 21101 y $2.00. Captura FAC-PRUEBA-0001.","Factura Devengada y póliza DEV automática.","19_factura_devengada.png")
note("Importes de la prueba:", "La cotización adjudicada y la OC son por $1.00. La suficiencia/compromiso usaron el promedio de tres cotizaciones ($2.00), por eso la factura vinculada al compromiso quedó por $2.00. En operación normal conviene que el compromiso definitivo se ajuste al monto adjudicado antes de facturar.","FFF2CC",(127,96,0))
step(16,"Recibir en almacén","Almacén > Recepción de Pedidos","Selecciona OC-2026-0001 y pulsa Registrar entrada. Confirma 1 unidad, costo $1.00 y área ALMACÉN TULTITLÁN.","Pendiente 0 y una entrada trazable.","20_entrada_almacen_completa.png")
step(17,"Comprobar pólizas","Contabilidad > Pólizas","Filtra/ubica pólizas 8, CP1 y DEV.","Las tres muestran Debe=Haber y diferencia $0.00.","21_polizas_presupuesto_compromiso_devengado.png")
page()

doc.add_heading("5. Qué hacer ante mensajes frecuentes",level=1)
table(["Mensaje","Qué significa","Acción exacta"],[
["El solicitante no está activo o no pertenece al área","La persona no está habilitada en el área","Configuración > Sistema > Usuarios; vincular persona, asignar área y marcar solicitante. Responsable: administrador del sistema."],
["No hay posición presupuestal válida","No existe saldo/clasificación para la partida","Presupuesto > Presupuesto autorizado; crear o ampliar posición 21101. Responsable: Presupuesto."],
["Falta matriz de conversión","No hay cuentas contables por etapa","Contabilidad > Matriz de conversión; configurar 21101 para programa, TG y año. Responsable: Contabilidad."],
["Debe existir compromiso vigente","La OC intenta avanzar sin AUT autorizado","Presupuesto > Estado de Contratos; crear y autorizar AUT. Responsable: Presupuesto/Compras."],
["Los detalles deben sumar el total","Factura sin partidas o monto inconsistente","Elegir el contrato; verificar que cargue sus partidas y que su suma sea igual al total."],
["Orden sin renglones por recibir","Orden no autorizada, ya surtida o año incorrecto","Revisar estado Por surtir, ejercicio 2026 y cantidades pendientes."],
])
doc.add_heading("6. Evidencia contable final",level=1)
table(["Póliza","Evento","Debe","Haber","Diferencia"],[
["8","Presupuesto de egresos autorizado","$104,833.35","$104,833.35","$0.00"],
["CP1","Compromiso AUT-1","$2.00","$2.00","$0.00"],
["DEV-20260810-…","Devengado de FAC-PRUEBA-0001","$2.00","$2.00","$0.00"],
])
doc.add_heading("7. Criterios de aceptación",level=1)
for txt in [
"REQ-2026-000002 usa la partida 21101, no Honorarios.",
"Existe una suficiencia autorizada y un compromiso AUT-1 vigente.",
"OC-2026-0001 está autorizada y en estado Por surtir antes de recibir.",
"FAC-PRUEBA-0001 tiene detalle que suma el total y genera póliza DEV.",
"La entrada de almacén está ligada al detalle de la OC y deja pendiente cero.",
"Las pólizas de autorizado, comprometido y devengado están balanceadas."
]:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(txt)
note("Conclusión:", "La prueba integral fue satisfactoria después de corregir los enlaces de captura, la carga de partidas de factura, las transacciones resilientes y la visibilidad de la acción de almacén.","E2F0D9",(46,92,42))

doc.save(DOCX)
print(DOCX)
