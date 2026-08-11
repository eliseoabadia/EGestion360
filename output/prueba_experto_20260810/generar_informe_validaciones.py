from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

ROOT = Path(__file__).resolve().parent
OLD = ROOT.parent / "prueba_funcional_20260810" / "capturas"
NEW = ROOT / "capturas"
OUT = ROOT / "Manual_visual_prueba_integral_y_validaciones.docx"

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(.55); sec.bottom_margin = Inches(.55)
sec.left_margin = Inches(.62); sec.right_margin = Inches(.62)
styles = doc.styles
styles['Normal'].font.name = 'Aptos'; styles['Normal'].font.size = Pt(9)
styles['Title'].font.name = 'Aptos Display'; styles['Title'].font.size = Pt(27)
styles['Title'].font.color.rgb = RGBColor(24, 71, 120)
for s in ('Heading 1','Heading 2'):
    styles[s].font.name='Aptos Display'; styles[s].font.color.rgb=RGBColor(24,71,120)

def title(text, subtitle=None):
    p=doc.add_paragraph(); p.style='Title'; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(text)
    if subtitle:
        q=doc.add_paragraph(subtitle); q.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def shot(path, caption):
    path=Path(path)
    if path.exists():
        doc.add_picture(str(path), width=Inches(7.0))
        p=doc.paragraphs[-1]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        c=doc.add_paragraph(caption); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
        c.runs[0].italic=True; c.runs[0].font.size=Pt(8)

def bullets(items):
    for x in items: doc.add_paragraph(x, style='List Bullet')

title('Manual visual de prueba integral', 'Presupuesto → compra → factura → almacén → pólizas | Validaciones para usuarios inexpertos | 10 de agosto de 2026')
doc.add_heading('Resultado ejecutivo',1)
doc.add_paragraph('Se limpió el flujo transaccional conservando catálogos, usuarios y bienes patrimoniales. Se repitió la captura presupuestal y se probaron errores de usuario. Los rechazos de mutaciones HTTP quedan en SIS.SystemLog con usuario, ruta, estado, duración y TraceId, sin almacenar contraseñas, tokens ni cuerpos de petición.')
bullets([
 'Corregido: un anteproyecto ya no puede guardarse con total cero ni con meses negativos, ni por pantalla ni por API.',
 'Corregido: los importes mensuales recalculan el total en tiempo real.',
 'Corregido: el selector de anteproyecto usa una altura estable; evitaba una rejilla visualmente vacía.',
 'Añadido: una recepción parcial genera notificación interna al creador de la orden con recibido y pendiente.',
 'Detectado durante la repetición: la autorización presupuestal del servidor actual rechaza el alta porque el procedimiento desplegado intenta insertar FKIdEmpresa_SIS nulo. El intento quedó en bitácora; no se insertó un registro inconsistente.'
])

doc.add_heading('1. Preparación y limpieza segura',1)
doc.add_paragraph('Ejecutar Scripts/20260810_LimpiarFlujoPruebaIntegral.sql únicamente después de confirmar respaldo. El script valida el nombre de la base, usa transacción y elimina sólo tablas transaccionales del flujo; conserva catálogos maestros.')

doc.add_heading('2. Validaciones iniciales',1)
shot(NEW/'01_anteproyecto_campos_obligatorios.png','Campos obligatorios: el sistema indica qué debe completar el usuario.')
shot(NEW/'03_anteproyecto_total_cero_bloqueado.png','Total $0.00 bloqueado con instrucción concreta.')
shot(NEW/'04_anteproyecto_valido_12000.png','Captura válida: $12,000 en agosto y clasificación presupuestal completa.')

doc.add_heading('3. Camino operativo completo',1)
doc.add_paragraph('Las siguientes evidencias visuales corresponden al recorrido funcional completo ejecutado en el mismo entorno antes de la limpieza. Sirven como guía de operación; la repetición actual se detuvo de forma segura en autorización por el defecto de base descrito en la sección 7.')
steps=[
 ('03_requisicion_creada.png','1. Crear requisición.'),
 ('07_requisicion_completa.png','2. Completar partidas y bienes; la suma debe coincidir con el importe.'),
 ('11_tres_cotizaciones.png','3. Registrar y comparar cotizaciones de proveedores.'),
 ('14_suficiencia_autorizada.png','4. Solicitar y autorizar suficiencia presupuestal.'),
 ('18_orden_compra_autorizada.png','5. Autorizar la orden de compra.'),
 ('19_factura_devengada.png','6. Registrar factura y devengado.'),
 ('20_entrada_almacen_completa.png','7. Recibir en almacén.'),
 ('21_polizas_presupuesto_compromiso_devengado.png','8. Confirmar pólizas de autorizado, comprometido y devengado.')]
for f,c in steps: shot(OLD/f,c)

doc.add_heading('4. Matriz de errores que debe enfrentar el sistema',1)
table=doc.add_table(rows=1, cols=4); table.style='Light Shading Accent 1'
for i,h in enumerate(['Caso','Respuesta esperada','Persistencia','Aviso']): table.rows[0].cells[i].text=h
rows=[
('Fecha final anterior a inicial','Bloquear y explicar el orden correcto','Bitácora WARN del rechazo','Snackbar; sin correo'),
('Importe cero o negativo','Bloquear antes de guardar','Bitácora si llega a API','Snackbar'),
('Importe superior al saldo','Bloquear y mostrar saldo disponible','Bitácora WARN','Notificación a Presupuesto si requiere ampliación'),
('Factura con fecha previa al contrato','Bloquear','Bitácora WARN','Mensaje en formulario'),
('Factura cuyo detalle no suma el total','Bloquear y mostrar diferencia','Bitácora WARN','Mensaje en formulario'),
('Entrega mayor al pendiente','Bloquear OVER_RECEIPT','Bitácora WARN','Mensaje a Almacén'),
('Entrega parcial','Aceptar; conservar saldo pendiente','Movimiento + auditoría normal','Notificación interna al comprador'),
('Entrega completa','Aceptar y cerrar pendiente','Movimiento de almacén','Notificación opcional de cumplimiento')]
for row in rows:
    cells=table.add_row().cells
    for i,v in enumerate(row): cells[i].text=v

doc.add_heading('5. Política recomendada de mensajes y correo',1)
bullets([
 'No enviar correo por errores de captura corregibles; usar mensaje inmediato y bitácora.',
 'Enviar notificación interna por recepción parcial, exceso rechazado, suficiencia detenida y factura incompleta.',
 'Usar correo sólo para pendientes que requieren a otra área o exceden un SLA; incluir liga al expediente, responsable y acción esperada.',
 'Nunca enviar contraseñas, tokens, cuerpos completos ni datos fiscales sensibles en la bitácora o correo.',
 'No aceptar entregas de más automáticamente: bloquear, levantar incidencia y exigir ampliación/modificación de orden o devolución documentada.'
])

doc.add_heading('6. Qué hacer ante entregas incompletas o excedentes',1)
doc.add_paragraph('Entrega incompleta: registrar únicamente lo físicamente recibido, mantener la orden con saldo pendiente y notificar al comprador. La factura debe devengarse sólo por lo aceptado conforme a la política institucional. Entrega excedente: no aumentar existencias; separar físicamente el excedente y decidir devolución o modificación formal de la orden antes de registrar.')

doc.add_heading('7. Hallazgo bloqueante de la repetición',1)
doc.add_paragraph('La API POST /api/EgresoAutorizado/autorizar-proyectado/1 devolvió HTTP 400: el procedimiento desplegado intentó insertar NULL en PRES.EgresoAutorizado.FKIdEmpresa_SIS. La bitácora registró usuario, ruta y TraceId. La transacción evitó presupuesto autorizado y póliza huérfanos. Antes de continuar con datos nuevos debe corregirse y desplegarse el procedimiento, repetir la autorización y verificar que se generen EgresoAutorizado y la póliza balanceada en una sola transacción.')

doc.add_heading('8. Lista de verificación final',1)
bullets([
 'Anteproyecto total > 0 y clasificación completa.',
 'Presupuesto autorizado y póliza balanceada.',
 'Requisición dentro del saldo; partidas = importe.',
 'Cotizaciones con fechas coherentes y proveedor activo.',
 'Suficiencia autorizada antes de la orden.',
 'Orden autorizada; contrato cuando aplique.',
 'Factura: fecha válida y detalles = total.',
 'Recepción: nunca mayor al pendiente; parcial notificada.',
 'Pólizas de autorizado, comprometido y devengado visibles y balanceadas.',
 'SIS.SystemLog contiene rechazos sin datos sensibles.'
])

for section in doc.sections:
    footer=section.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('EGestion360 · Prueba integral y validaciones · 2026')
doc.save(OUT)
print(OUT)
