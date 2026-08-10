from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
CAP = ROOT / "capturas"
OUT = ROOT / "Manual_visual_presupuesto_a_orden_compra.docx"

BLUE = "2E74B5"
NAVY = "17365D"
LIGHT = "EAF2F8"
GREEN = "E2F0D9"
AMBER = "FFF2CC"
GRAY = "F2F2F2"
RED = "FCE4D6"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.72)
sec.bottom_margin = Inches(0.68)
sec.left_margin = Inches(0.75)
sec.right_margin = Inches(0.75)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.12
for name, size, color in [("Title", 25, NAVY), ("Heading 1", 16, BLUE), ("Heading 2", 13, NAVY), ("Heading 3", 11.5, NAVY)]:
    st = styles[name]
    st.font.name = "Calibri"
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = True
    st.paragraph_format.space_before = Pt(8)
    st.paragraph_format.space_after = Pt(5)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_text(cell, text, bold=False, color=None, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def info_box(title, body, fill=LIGHT):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.85)
    c = table.cell(0, 0)
    shade(c, fill)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    r.font.size = Pt(10.5)
    r2 = p.add_run(body)
    r2.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

def route(path, owner):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(4.85)
    table.columns[1].width = Inches(2.0)
    set_cell_text(table.cell(0,0), "MENÚ  ›  " + path, True, NAVY, 9)
    set_cell_text(table.cell(0,1), "RESPONSABLE  ›  " + owner, True, NAVY, 9)
    shade(table.cell(0,0), LIGHT); shade(table.cell(0,1), LIGHT)

def steps(items):
    for i, text in enumerate(items, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(text)
        p.paragraph_format.space_after = Pt(3)

def screenshot(filename, caption):
    path = CAP / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.85))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(3)
    c.runs[0].italic = True
    c.runs[0].font.size = Pt(8.5)
    c.runs[0].font.color.rgb = RGBColor(89,89,89)

def new_page(title=None, subtitle=None):
    doc.add_page_break()
    if title:
        doc.add_heading(title, level=1)
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
        p.runs[0].bold = True

def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("PCI · Manual visual · ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)

add_footer(sec)

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(64)
r = p.add_run("PCI")
r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor.from_string(BLUE)
p = doc.add_paragraph(style="Title")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Del presupuesto proyectado\na la orden de compra")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Manual visual paso a paso · Ejercicio completo en EGestion360")
r.font.size = Pt(13); r.font.color.rgb = RGBColor.from_string(BLUE)
doc.add_paragraph()
info_box("CASO PRÁCTICO", "Compra de 10 millares de papel bond tamaño carta para ALMACÉN TULTITLÁN. Presupuesto: $12,000. Requisición: $2,500. Cotización adjudicada y orden: $2,400.", GREEN)
table = doc.add_table(rows=5, cols=2)
data = [
    ("Usuario", "ADMIN001 · Gabriela Corona Espinosa"),
    ("Entidad", "Plataforma de Compras Integral · Sucursal Operativa"),
    ("Área", "208C0101320202T · ALMACÉN TULTITLÁN"),
    ("Ejercicio", "2026"),
    ("Fecha de ejecución", "10 de agosto de 2026"),
]
for i,(a,b) in enumerate(data):
    set_cell_text(table.cell(i,0), a, True, NAVY)
    set_cell_text(table.cell(i,1), b)
    shade(table.cell(i,0), LIGHT)
doc.add_paragraph()
p = doc.add_paragraph("Documento generado sobre una base local/de prueba. No contiene contraseñas. Antes de repetir el ejercicio en producción, valide permisos, ejercicio, entidad y saldos.")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].italic = True; p.runs[0].font.size = Pt(9)

# Overview
new_page("1. Ruta completa y resultado")
doc.add_paragraph("El expediente se construye en este orden. No conviene saltar etapas porque la aplicación valida dependencias entre usuario, área, presupuesto, cotizaciones y proveedor adjudicado.")
flow = doc.add_table(rows=8, cols=3)
flow.style = "Table Grid"
headers=("Etapa","Resultado del caso","Menú principal")
for j,h in enumerate(headers): set_cell_text(flow.cell(0,j),h,True,"FFFFFF"); shade(flow.cell(0,j),BLUE)
rows=[
    ("0. Preparar usuario","Solicitante habilitada en ALMACÉN TULTITLÁN","Configuración › Sistema › Usuario"),
    ("1. Proyectar","Anteproyecto 2 · $12,000 en agosto","Presupuesto › Egreso › Proyectado"),
    ("2. Autorizar","Presupuesto autorizado 12","Presupuesto › Egreso › Autorizado"),
    ("3. Requerir","REQ-000011 · $2,500","Adquisiciones › Requisición"),
    ("4. Cotizar","Cotizaciones 26, 27 y 28","Adquisiciones › Cotización"),
    ("5. Suficiencia","Solicitud 15 · autorización 13","Presupuesto › Comprometido › Suficiencia"),
    ("6. Comprometer y ordenar","COMP-MANUAL-2026-001 · OC-2026-0001","Presupuesto › Comprometido / Adquisiciones › Orden de Compra"),
]
for i,row in enumerate(rows,1):
    for j,v in enumerate(row): set_cell_text(flow.cell(i,j),v,size=8.6)
    if i%2==0:
        for j in range(3): shade(flow.cell(i,j),GRAY)
doc.add_heading("Datos de control", level=2)
info_box("IMPORTES", "Presupuesto disponible $12,000 → requisición/suficiencia $2,500 → oferta ganadora $2,400 → orden $2,400. El compromiso conserva $2,500 porque corresponde al techo autorizado; la orden usa el precio adjudicado.", AMBER)
info_box("DEFINICIÓN DE ÉXITO", "La orden se considera completa cuando muestra 1 detalle, 1 partida, total $2,400 y el proveedor GRUPO EMPRESARIAL EMPROVE.", GREEN)

# User
new_page("2. Preparación: usuario, persona y área", "Paso previo obligatorio")
route("Configuración › Sistema › Usuario", "Administrador del sistema")
steps([
    "Busque a Gabriela Corona Espinosa y abra Editar usuario.",
    "En Áreas asignadas seleccione ALMACÉN TULTITLÁN.",
    "Active Adscrito y Es solicitante; asigne el área y confirme que el contador aumente.",
    "Conserve activo al usuario. El departamento no sustituye el permiso de solicitante por área."
])
screenshot("04_usuario_area_asignada.png", "Resultado esperado: el área aparece asignada y marcada como solicitante.")
info_box("SI APARECE ‘EL SOLICITANTE NO ESTÁ ACTIVO O NO PERTENECE AL ÁREA’", "Vuelva a esta pantalla. Verifique usuario activo, persona vinculada, área correcta y la marca Es solicitante. Responsable: administrador del sistema o personal.", RED)

# Projected
new_page("3. Crear el presupuesto proyectado")
route("Presupuesto › Egreso › Proyectado › Anteproyecto", "Área de Presupuesto")
steps([
    "Pulse Nuevo y seleccione empresa, ejercicio 2026, programa 02030201 y área ALMACÉN TULTITLÁN.",
    "Seleccione partida 21101, Ingresos propios, TG 1, DI 1, DG 1 y proyecto 10766.",
    "Capture la descripción del caso y $12,000 en agosto. El mes debe coincidir con la fecha posterior de la requisición.",
    "Guarde y confirme que el total sea $12,000."
])
screenshot("05_anteproyecto_capturado.png", "Ejemplo inicial: la imagen muestra septiembre; para este caso cambie el importe a agosto antes de guardar.")
info_box("CONTROL CLAVE", "Si la requisición es de agosto, la suficiencia debe encontrar saldo en agosto. Un saldo colocado sólo en septiembre no puede respaldar un gasto de agosto.", AMBER)

# Authorization
new_page("4. Autorizar el presupuesto")
route("Presupuesto › Egreso › Proyectado › Autorizar", "Responsable presupuestal")
steps([
    "Localice el anteproyecto creado y pulse Autorizar.",
    "Revise programa, partida, fuente, área, proyecto y total.",
    "Confirme la autorización y verifique el registro en Presupuesto autorizado."
])
screenshot("08_autorizar_anteproyecto.png", "Confirmación previa: no autorice si el importe o la clasificación no corresponden.")

# Requisition
new_page("5. Crear la requisición y sus partidas")
route("Adquisiciones › Requisición", "Área solicitante / Compras")
steps([
    "Pulse Nuevo. Capture $2,500, tipo Bien, procedimiento Ordinario y solicitante Gabriela Corona Espinosa.",
    "Elija la posición presupuestal autorizada. La aplicación completa programa, fuente, TG, DI y DG.",
    "Guarde. El caso crea REQ-000011.",
    "Abra la requisición, agregue la partida 21101 por $2,500 y después el bien MO00394: 10 MILLARES de papel bond tamaño carta."
])
screenshot("12_requisicion_formulario.png", "Formulario de requisición con clasificación presupuestal.")
screenshot("16_requisicion_completa.png", "Resultado: requisición con partida y detalle listos para cotizar.")
info_box("SI EL BIEN NO SE PUEDE AGREGAR", "Revise que esté activo en Almacén, tenga unidad de medida y disponibilidad/configuración aplicable. Responsable: Almacén; para clasificación, Presupuesto/Contabilidad.", RED)

# Quotes
new_page("6. Registrar tres cotizaciones")
route("Adquisiciones › Cotización", "Compras / Padrón de proveedores")
steps([
    "Cree una cotización para la requisición 11 y seleccione un proveedor activo.",
    "Agregue el mismo bien cotizado y capture su precio unitario.",
    "Repita con proveedores distintos hasta completar tres ofertas.",
    "Compare totales y vigencia. Para este caso se adjudica COT-27 por $2,400."
])
screenshot("17_cotizacion_formulario.png", "Alta de cotización vinculada a la requisición.")
screenshot("21_tres_cotizaciones.png", "Resultado: tres cotizaciones completas de proveedores diferentes.")
q = doc.add_table(rows=4, cols=3); q.style="Table Grid"
for j,h in enumerate(("Cotización","Proveedor","Total")): set_cell_text(q.cell(0,j),h,True,"FFFFFF"); shade(q.cell(0,j),BLUE)
for i,row in enumerate((("26","Proveedor 1","$2,500"),("27","GRUPO EMPRESARIAL EMPROVE","$2,400 · adjudicada"),("28","BFS INGENIERÍA APLICADA","$2,600")),1):
    for j,v in enumerate(row): set_cell_text(q.cell(i,j),v)
    if i==2:
        for j in range(3): shade(q.cell(i,j),GREEN)

# Sufficiency
new_page("7. Solicitar y autorizar suficiencia")
route("Presupuesto › Egreso › Comprometido › Solicitud de Suficiencia", "Área solicitante / Presupuesto")
steps([
    "Seleccione la requisición 11 y ejecute el precálculo.",
    "Confirme que se detecten tres cotizaciones completas, partida 21101, mes agosto y total $2,500.",
    "Cree la solicitud; el caso genera la solicitud 15.",
    "En Autorización de Suficiencia seleccione la solicitud, a la persona autorizadora y cambie el estatus a Autorizada. El caso genera autorización 13."
])
screenshot("22_suficiencia_precalculo.png", "Precálculo: revise antes de generar la solicitud.")
screenshot("24_autorizacion_suficiencia.png", "Autorización: use una persona con atribución presupuestal.")
info_box("NO CONTINÚE SI…", "Falta una cotización completa, el mes no tiene saldo, la clasificación difiere o la persona no puede autorizar. Corrija en la etapa de origen y vuelva a calcular.", RED)

# Commitment
new_page("8. Registrar el compromiso presupuestal")
route("Presupuesto › Egreso › Comprometido › Registro Comprometido", "Presupuesto / Contratos")
steps([
    "Pulse Nuevo y seleccione la autorización de suficiencia 13.",
    "Seleccione el mismo proveedor de la cotización adjudicada: GRUPO EMPRESARIAL EMPROVE.",
    "Capture número COMP-MANUAL-2026-001, fecha 10/08/2026, vigencia y monto autorizado $2,500.",
    "Guarde con estatus Vigente."
])
screenshot("29_registro_comprometido_creado.png", "Resultado: compromiso vigente por el techo autorizado de $2,500.")
info_box("POR QUÉ SON $2,500 Y NO $2,400", "El compromiso se vincula con la autorización de suficiencia y debe corresponder a su importe. La orden sí toma el precio adjudicado de $2,400. Esta diferencia representa un remanente presupuestal de $100.", AMBER)

# Order
new_page("9. Generar y completar la orden de compra")
route("Adquisiciones › Orden de Compra", "Compras")
steps([
    "Pulse Nuevo, seleccione la requisición 11 y la cotización COT-27.",
    "Compruebe que el proveedor se complete automáticamente y cree la orden.",
    "Abra OC-2026-0001. En Detalles pulse Agregar y seleccione MO00394; se cargan 10, MILLAR y $240.",
    "En Partidas presupuestales pulse Agregar, seleccione 21101, Ingresos propios e importe $2,400.",
    "Verifique total $2,400, 1 detalle y 1 partida."
])
screenshot("30_orden_compra_final_formulario.png", "La cotización adjudicada determina el proveedor y el precio de la orden.")
screenshot("31_orden_compra_creada_final.png", "Resultado final: OC-2026-0001 completa por $2,400.")

# Troubleshooting
new_page("10. Guía rápida de bloqueos")
issues = [
    ("Solicitante sin área o permiso", "Configuración › Sistema › Usuario", "Administrador del sistema", "Asignar área y marcar Es solicitante."),
    ("No hay posición presupuestal", "Presupuesto › Egreso › Proyectado/Autorizado", "Presupuesto", "Crear y autorizar saldo en el mismo mes de la requisición."),
    ("No permite cotizar", "Adquisiciones › Requisición/Cotización", "Compras", "Completar partida, bien y proveedor activo."),
    ("Exige tres cotizaciones", "Adquisiciones › Cotización", "Compras", "Registrar tres proveedores distintos con precio para todos los bienes."),
    ("Suficiencia insuficiente", "Presupuesto › Comprometido › Suficiencia", "Presupuesto", "Revisar mes, partida, fuente e importe."),
    ("COMMITMENT_REQUIRED", "Presupuesto › Comprometido › Registro Comprometido", "Presupuesto/Contratos", "Crear compromiso vigente con la misma autorización y proveedor adjudicado."),
    ("Orden con $0 o sin hijos", "Adquisiciones › Orden de Compra › Ver detalle", "Compras", "Agregar Detalles y Partidas; no basta con crear el encabezado."),
]
t = doc.add_table(rows=1+len(issues), cols=4); t.style="Table Grid"
for j,h in enumerate(("Mensaje/síntoma","Dónde ir","Responsable","Acción")): set_cell_text(t.cell(0,j),h,True,"FFFFFF",8.5); shade(t.cell(0,j),BLUE)
for i,row in enumerate(issues,1):
    for j,v in enumerate(row): set_cell_text(t.cell(i,j),v,size=8.0)
    if i%2==0:
        for j in range(4): shade(t.cell(i,j),GRAY)
doc.add_heading("Checklist antes de terminar", level=2)
for text in [
    "Usuario activo, persona vinculada, área asignada y permiso solicitante.",
    "Presupuesto autorizado en el mes correcto y por importe suficiente.",
    "Requisición con partida y bien; tres cotizaciones completas.",
    "Suficiencia autorizada y compromiso vigente con el proveedor adjudicado.",
    "Orden con detalle, partida e importes coincidentes."
]:
    p=doc.add_paragraph(style="List Bullet"); p.add_run("☐ "+text)

# Audit
new_page("11. Trazabilidad del ejercicio")
records = [
    ("Anteproyecto", "2", "$12,000", "Agosto 2026"),
    ("Presupuesto autorizado", "12", "$12,000", "Partida 21101"),
    ("Requisición", "11 / REQ-000011", "$2,500", "10 millares"),
    ("Cotizaciones", "26, 27, 28", "$2,500 / $2,400 / $2,600", "Ganadora: 27"),
    ("Solicitud suficiencia", "15", "$2,500", "Autorización 13"),
    ("Compromiso", "14 / COMP-MANUAL-2026-001", "$2,500", "Vigente"),
    ("Orden de compra", "6 / OC-2026-0001", "$2,400", "Inicial, completa"),
]
t=doc.add_table(rows=1+len(records),cols=4); t.style="Table Grid"
for j,h in enumerate(("Entidad","Identificador","Importe","Referencia")): set_cell_text(t.cell(0,j),h,True,"FFFFFF"); shade(t.cell(0,j),BLUE)
for i,row in enumerate(records,1):
    for j,v in enumerate(row): set_cell_text(t.cell(i,j),v,size=8.8)
    if i%2==0:
        for j in range(4): shade(t.cell(i,j),GRAY)
doc.add_paragraph()
info_box("CIERRE DEL CASO", "La compra quedó trazada desde el recurso proyectado hasta OC-2026-0001. La orden está en estatus INICIAL, con 1 detalle y 1 partida; el siguiente paso operativo sería revisar y autorizar la orden según los permisos y políticas de la entidad.", GREEN)
p=doc.add_paragraph("Fin del manual · Conserve este documento junto con las evidencias del expediente.")
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold=True; p.runs[0].font.color.rgb=RGBColor.from_string(NAVY)

doc.save(OUT)
print(OUT)
