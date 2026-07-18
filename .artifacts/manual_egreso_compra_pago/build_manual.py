from __future__ import annotations

import os
import sys
from datetime import date
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

SKILL_ROOT = Path(r"C:\Users\Eliseo\.codex\plugins\cache\openai-primary-runtime\documents\26.630.12135\skills\documents")
sys.path.insert(0, str(SKILL_ROOT / "scripts"))
from table_geometry import apply_table_geometry  # type: ignore


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "Manual_Flujo_Anteproyecto_Egreso_a_Cheque_y_Entrada_Bienes.docx"
FLOW_PATH = OUT_DIR / "flujo_integral.png"

NAVY = "17365D"
BLUE = "2E74B5"
BLUE_DARK = "1F4D78"
PURPLE = "65578F"
TEAL = "2C7468"
GOLD = "B7791F"
RED = "9B1C1C"
GREEN = "26734D"
INK = "1F2937"
MUTED = "667085"
PALE_BLUE = "E8EEF5"
PALE_GRAY = "F2F4F7"
PALE_GOLD = "FFF4D6"
PALE_RED = "FDECEC"
PALE_GREEN = "E8F4EE"
WHITE = "FFFFFF"
BLACK = "000000"

CONTENT_DXA = 9360
TABLE_INDENT = 120
CELL_MARGINS = {"top": 90, "bottom": 90, "start": 120, "end": 120}


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = "D0D5DD", size: int = 6) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def set_cell_text(cell, text: str, bold: bool = False, color: str = INK, size: float = 9.2,
                  align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    r._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    r.font.size = Pt(size)
    r.font.color.rgb = rgb(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int],
              header_fill: str = PALE_BLUE, font_size: float = 9.1,
              alignments: list | None = None) -> object:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        set_cell_text(hdr.cells[i], header, bold=True, color=NAVY, size=9.0,
                      align=(alignments[i] if alignments else WD_ALIGN_PARAGRAPH.LEFT))
        set_cell_shading(hdr.cells[i], header_fill)
        set_cell_borders(hdr.cells[i])
    for row_data in rows:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for i, value in enumerate(row_data):
            set_cell_text(row.cells[i], value, size=font_size,
                          align=(alignments[i] if alignments else WD_ALIGN_PARAGRAPH.LEFT))
            set_cell_borders(row.cells[i])
    apply_table_geometry(table, widths, table_width_dxa=sum(widths), indent_dxa=TABLE_INDENT,
                         cell_margins_dxa=CELL_MARGINS)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_label_detail_table(doc: Document, pairs: list[tuple[str, str]]) -> object:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    set_cell_text(header.cells[0], "Referencia", bold=True, color=NAVY, size=8.8)
    set_cell_text(header.cells[1], "Detalle", bold=True, color=NAVY, size=8.8)
    for cell in header.cells:
        set_cell_shading(cell, PALE_BLUE)
        set_cell_borders(cell, color="D8DEE8", size=5)
    for label, value in pairs:
        row = table.add_row()
        set_cell_text(row.cells[0], label, bold=True, color=NAVY, size=9.2)
        set_cell_shading(row.cells[0], PALE_GRAY)
        set_cell_text(row.cells[1], value, color=INK, size=9.2)
        for cell in row.cells:
            set_cell_borders(cell, color="D8DEE8", size=5)
    apply_table_geometry(table, [2100, 7260], table_width_dxa=CONTENT_DXA,
                         indent_dxa=TABLE_INDENT, cell_margins_dxa=CELL_MARGINS)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, sep, text, end])


def add_page_number_footer(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("eGestion360  |  Manual operativo  |  Página ")
    r.font.name = "Calibri"
    r.font.size = Pt(8.5)
    r.font.color.rgb = rgb(MUTED)
    add_field(p, "PAGE")
    r = p.add_run(" de ")
    r.font.size = Pt(8.5)
    r.font.color.rgb = rgb(MUTED)
    add_field(p, "NUMPAGES")


def add_running_header(section) -> None:
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("GESTIÓN EMPRESARIAL 360  /  EGRESO · COMPRA · RECEPCIÓN · PAGO")
    r.font.name = "Calibri"
    r.font.size = Pt(8.2)
    r.bold = True
    r.font.color.rgb = rgb(MUTED)


def ensure_paragraph_border(paragraph, side: str, color: str, size: int = 14, space: int = 6) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)
    p_bdr.append(edge)


def add_callout(doc: Document, label: str, text: str, kind: str = "info") -> None:
    palette = {
        "info": (PALE_BLUE, BLUE),
        "control": (PALE_GREEN, GREEN),
        "warning": (PALE_GOLD, GOLD),
        "risk": (PALE_RED, RED),
    }
    fill, accent = palette[kind]
    p = doc.add_paragraph(style="Callout")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.10)
    p.paragraph_format.right_indent = Inches(0.06)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    ensure_paragraph_border(p, "left", accent, size=22, space=7)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.color.rgb = rgb(accent)
    r2 = p.add_run(text)
    r2.font.color.rgb = rgb(INK)


def define_numbering(doc: Document, fmt: str, text_value: str, left: int, hanging: int,
                     color: str = NAVY) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text_value)
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    col = OxmlElement("w:color")
    col.set(qn("w:val"), color)
    r_pr.append(col)
    lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbered_steps(doc: Document, steps: list[str]) -> None:
    num_id = define_numbering(doc, "decimal", "%1.", 540, 270)
    for step in steps:
        p = doc.add_paragraph(style="Step")
        num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_el)
        p.add_run(step)


def add_bullets(doc: Document, items: list[str]) -> None:
    num_id = define_numbering(doc, "bullet", "•", 540, 270, color=BLUE_DARK)
    for item in items:
        p = doc.add_paragraph(style="Bullet")
        num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_el = OxmlElement("w:numId")
        num_id_el.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_el)
        p.add_run(item)


def add_checklist(doc: Document, items: list[str]) -> None:
    rows = [["☐", item] for item in items]
    table = add_table(doc, ["", "Verificación"], rows, [520, 8840], header_fill=PALE_GRAY,
                      font_size=9.3, alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])
    return table


def add_heading(doc: Document, text: str, level: int = 1) -> object:
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_body(doc: Document, text: str, bold_lead: str | None = None) -> object:
    p = doc.add_paragraph(style="Normal")
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        r1.bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_stage(doc: Document, number: int, title: str, route: str, owner: str, entry: str,
              objective: str, steps: list[str], controls: list[str], output: str,
              correction: str, fields: str | None = None, auth: str | None = None) -> None:
    add_heading(doc, f"{number}. {title}", 2)
    add_label_detail_table(doc, [
        ("Ruta", route),
        ("Responsable principal", owner),
        ("Entrada mínima", entry),
        ("Resultado", output),
    ])
    add_body(doc, objective)
    if fields:
        add_body(doc, f"Campos clave. {fields}")
    add_heading(doc, "Procedimiento", 3)
    add_numbered_steps(doc, steps)
    if auth:
        add_callout(doc, "Solicitud / autorización", auth, "control")
    add_heading(doc, "Controles antes de avanzar", 3)
    add_bullets(doc, controls)
    add_callout(doc, "Si requiere corrección", correction, "warning")


def find_font(size: int, bold: bool = False):
    candidates = [
        Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / ("calibrib.ttf" if bold else "calibri.ttf"),
        Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / ("arialbd.ttf" if bold else "arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def rounded_box(draw, xy, fill, outline, radius=22, width=4):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def wrap_text(draw, text: str, font, max_width: int) -> list[str]:
    words = text.split()
    lines, current = [], ""
    for word in words:
        candidate = word if not current else current + " " + word
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def create_flow_image(path: Path) -> None:
    width, height = 2400, 1280
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = find_font(54, True)
    sub_font = find_font(28, False)
    box_font = find_font(31, True)
    small_font = find_font(24, False)
    num_font = find_font(28, True)
    draw.text((100, 60), "Flujo integral de egreso, adquisición, recepción y pago", font=title_font, fill="#17365D")
    draw.text((100, 130), "Los colores identifican el área líder; las flechas indican la dependencia mínima.", font=sub_font, fill="#667085")

    stages = [
        (1, "Anteproyecto\nde egresos", "Presupuesto"),
        (2, "Presupuesto\nautorizado", "Presupuesto"),
        (3, "Requisición", "Adquisiciones"),
        (4, "Cotización", "Adquisiciones"),
        (5, "Solicitud de\nsuficiencia", "Adquisiciones"),
        (6, "Autorización de\nsuficiencia", "Presupuesto"),
        (7, "Orden de\ncompra", "Adquisiciones"),
        (8, "Recepción / entrada\nde almacén", "Almacén"),
        (9, "Compromiso /\ncontrato", "Presupuesto"),
        (10, "Factura", "Tesorería"),
        (11, "CLC / provisión\nde pago", "Tesorería"),
        (12, "Cheque o\ntransferencia", "Tesorería"),
    ]
    colors = {
        "Presupuesto": ("#E8EEF5", "#2E74B5"),
        "Adquisiciones": ("#F0ECF7", "#65578F"),
        "Almacén": ("#E8F4EE", "#2C7468"),
        "Tesorería": ("#FFF4D6", "#B7791F"),
    }
    positions = []
    box_w, box_h = 500, 210
    xs = [100, 700, 1300, 1900]
    ys = [250, 560, 870]
    order_positions = [(xs[0], ys[0]), (xs[1], ys[0]), (xs[2], ys[0]), (xs[3], ys[0]),
                       (xs[3], ys[1]), (xs[2], ys[1]), (xs[1], ys[1]), (xs[0], ys[1]),
                       (xs[0], ys[2]), (xs[1], ys[2]), (xs[2], ys[2]), (xs[3], ys[2])]
    for x, y in order_positions:
        positions.append((x, y, x + box_w, y + box_h))

    for i in range(len(positions) - 1):
        x1, y1, x2, y2 = positions[i]
        nx1, ny1, nx2, ny2 = positions[i + 1]
        if abs(ny1 - y1) < 20:
            start = (x2 + 15, (y1 + y2) // 2)
            end = (nx1 - 15, (ny1 + ny2) // 2)
        elif nx1 < x1:
            start = ((x1 + x2) // 2, y2 + 8)
            end = ((nx1 + nx2) // 2, ny1 - 8)
        else:
            start = ((x1 + x2) // 2, y2 + 8)
            end = ((nx1 + nx2) // 2, ny1 - 8)
        draw.line([start, end], fill="#98A2B3", width=8)
        ex, ey = end
        if abs(end[0] - start[0]) > abs(end[1] - start[1]):
            direction = 1 if end[0] > start[0] else -1
            draw.polygon([(ex, ey), (ex - 24 * direction, ey - 15), (ex - 24 * direction, ey + 15)], fill="#98A2B3")
        else:
            direction = 1 if end[1] > start[1] else -1
            draw.polygon([(ex, ey), (ex - 15, ey - 24 * direction), (ex + 15, ey - 24 * direction)], fill="#98A2B3")

    for stage, pos in zip(stages, positions):
        n, label, area = stage
        fill, accent = colors[area]
        x1, y1, x2, y2 = pos
        rounded_box(draw, pos, fill, accent)
        draw.ellipse((x1 + 22, y1 + 22, x1 + 82, y1 + 82), fill=accent)
        nbox = draw.textbbox((0, 0), str(n), font=num_font)
        draw.text((x1 + 52 - (nbox[2] - nbox[0]) / 2, y1 + 52 - (nbox[3] - nbox[1]) / 2 - 2), str(n), font=num_font, fill="white")
        lines = label.split("\n")
        line_height = 40
        start_y = y1 + 45 if len(lines) == 2 else y1 + 66
        for idx, line in enumerate(lines):
            box = draw.textbbox((0, 0), line, font=box_font)
            tx = x1 + box_w / 2 - (box[2] - box[0]) / 2
            draw.text((tx, start_y + idx * line_height), line, font=box_font, fill="#1F2937")
        area_box = draw.textbbox((0, 0), area, font=small_font)
        draw.text((x1 + box_w / 2 - (area_box[2] - area_box[0]) / 2, y2 - 46), area, font=small_font, fill=accent)

    legend_y = 1180
    lx = 100
    for area, (fill, accent) in colors.items():
        draw.rounded_rectangle((lx, legend_y, lx + 44, legend_y + 28), radius=6, fill=fill, outline=accent, width=3)
        draw.text((lx + 58, legend_y - 4), area, font=small_font, fill="#475467")
        lx += 500
    img.save(path, quality=95)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(14)
    subtitle.font.color.rgb = rgb(BLUE_DARK)
    subtitle.paragraph_format.space_after = Pt(12)

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, BLUE_DARK, 10, 5),
    }
    for name, (size, color, before, after) in heading_specs.items():
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = rgb(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for custom_name in ("Step", "Bullet", "Callout"):
        if custom_name not in styles:
            styles.add_style(custom_name, 1)
        st = styles[custom_name]
        st.base_style = normal
        st.font.name = "Calibri"
        st.font.size = Pt(10.5 if custom_name != "Callout" else 10.2)
        st.font.color.rgb = rgb(INK)
        st.paragraph_format.space_before = Pt(0)
        st.paragraph_format.space_after = Pt(4 if custom_name != "Callout" else 6)
        st.paragraph_format.line_spacing = 1.25


def build_document() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    create_flow_image(FLOW_PATH)

    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True
    add_running_header(section)
    add_page_number_footer(section)

    # Cover: editorial_cover pattern with restrained corporate color.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(70)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MANUAL OPERATIVO")
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.bold = True
    r.font.color.rgb = rgb(GOLD)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Del anteproyecto de egresos\nal cheque y la entrada de bienes")
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Flujo completo de solicitud, autorización, adquisición, recepción y pago en eGestion360")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Presupuesto  ·  Adquisiciones  ·  Almacén  ·  Cuentas por pagar  ·  Tesorería")
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)
    r.font.color.rgb = rgb(MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(105)
    r = p.add_run("Versión 1.0  |  17 de julio de 2026")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = rgb(NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Documento elaborado con base en las rutas, validaciones y acciones disponibles en la solución actual.")
    r.font.size = Pt(9.5)
    r.font.italic = True
    r.font.color.rgb = rgb(MUTED)
    doc.add_page_break()

    add_heading(doc, "Control del documento", 1)
    add_label_detail_table(doc, [
        ("Propósito", "Guiar el ciclo completo desde la planeación del egreso hasta la recepción del bien y la emisión del cheque o transferencia."),
        ("Audiencia", "Planeación, áreas solicitantes, Compras, Presupuesto, Almacén, Cuentas por pagar, Tesorería y Control interno."),
        ("Alcance", "Flujo estándar de bienes y servicios; incluye solicitud y autorización de suficiencia, orden de compra, entrada, compromiso, factura, CLC y pago."),
        ("Sistema", "eGestion360 / módulos GRP."),
        ("Vigencia", "Validar nuevamente cuando cambien catálogos, permisos, rutas o reglas de negocio."),
    ])
    add_callout(doc, "Regla de lectura", "La descripción del estatus que muestra la pantalla es la fuente operativa. Los nombres pueden depender del catálogo configurado por la empresa.", "info")

    add_heading(doc, "Contenido", 1)
    add_numbered_steps(doc, [
        "Objetivo, alcance y criterios comunes.",
        "Roles y segregación de funciones.",
        "Mapa del flujo integral.",
        "Planeación: anteproyecto y presupuesto autorizado.",
        "Adquisición: requisición, cotización, suficiencia y orden de compra.",
        "Recepción: entrada de almacén y control de cantidades.",
        "Compromiso y pago: contrato, factura, CLC y cheque/transferencia.",
        "Controles, trazabilidad, correcciones y cierre del expediente.",
    ])

    add_heading(doc, "1. Objetivo y alcance", 1)
    add_body(doc, "Este manual explica el recorrido de un egreso planeado hasta su ejecución operativa y financiera. Presenta las entradas mínimas, responsables, acciones de pantalla, autorizaciones, validaciones y evidencias necesarias para conservar la trazabilidad de una compra.")
    add_body(doc, "El flujo estándar comienza con una partida mensual del anteproyecto, continúa con presupuesto autorizado, requisición y cotización, reserva suficiencia, formaliza la orden de compra, registra la recepción física y concluye con compromiso, factura, CLC y cheque o transferencia.")
    add_callout(doc, "Importante", "La orden de compra y la recepción física pertenecen a la rama operativa; el contrato, la factura, la CLC y el cheque forman la rama presupuestal-financiera. Ambas deben referirse a la misma empresa, requisición, partidas e importes.", "warning")

    add_heading(doc, "2. Roles y segregación de funciones", 1)
    add_table(doc,
              ["Rol", "Responsabilidad principal", "Acciones sensibles"],
              [
                  ["Planeación / Presupuesto", "Capturar el anteproyecto, revisar clasificación y convertirlo en presupuesto autorizado.", "Autorizar o regresar a proyectado."],
                  ["Área solicitante", "Definir necesidad, fechas, cantidades, partidas y justificación.", "Crear requisición y confirmar bienes/servicios."],
                  ["Compras / Adquisiciones", "Cotizar, seleccionar proveedor, solicitar suficiencia y preparar la orden.", "Emitir y autorizar orden de compra."],
                  ["Autorizador presupuestal", "Verificar saldo, clasificación, detalle e importe solicitado.", "Autorizar suficiencia y reservar presupuesto."],
                  ["Almacén", "Recibir contra orden, registrar cantidades, costo y documentos de entrega.", "Registrar entradas y vigilar pendientes."],
                  ["Cuentas por pagar", "Registrar contrato/factura, integrar partidas y generar CLC.", "Enviar factura a autorización."],
                  ["Tesorería", "Validar CLC, cuenta bancaria, póliza e importe.", "Generar cheque o transferencia."],
                  ["Control interno", "Revisar expediente y concordancia de importes.", "Consultar soportes y reportes; no sustituye al autorizador."],
              ],
              [1700, 4500, 3160], font_size=8.8)
    add_callout(doc, "Permisos", "Las acciones Crear, Editar, Eliminar y Autorizar se muestran según los permisos del usuario. Si falta una acción, solicite la asignación del permiso correspondiente; no duplique registros para evadir el control.", "control")

    add_heading(doc, "3. Mapa del flujo integral", 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(FLOW_PATH), width=Inches(6.45))
    picture_props = doc.inline_shapes[-1]._inline.docPr
    picture_props.set("descr", "Diagrama del flujo integral desde anteproyecto de egresos hasta cheque o transferencia, con orden de compra y entrada de almacén.")
    picture_props.set("title", "Flujo integral de egreso, adquisición, recepción y pago")
    p2 = doc.add_paragraph("Figura 1. Dependencias mínimas del flujo estándar.")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(8)
    for r in p2.runs:
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = rgb(MUTED)
    add_callout(doc, "Punto de control central", "Sin suficiencia autorizada, la requisición no debe generar una orden de compra. La autorización verifica además que exista detalle, clasificación completa y saldo disponible.", "risk")

    add_heading(doc, "4. Requisitos comunes antes de iniciar", 1)
    add_checklist(doc, [
        "Seleccionar la empresa o sucursal correcta y el año presupuestal vigente.",
        "Confirmar que programa, partida, área, fuente de financiamiento, tipo de gasto (TG), dígito identificador (DI) y destino del gasto (DG) estén activos.",
        "Verificar catálogos de bienes/servicios, unidades, proveedores, cuentas bancarias y pólizas.",
        "Contar con permisos para crear, actualizar y, cuando corresponda, autorizar.",
        "Preparar justificación, oficio, cotizaciones, orden, factura, remisión y demás soportes del expediente.",
        "Definir quién solicita, quién revisa y quién autoriza; evitar que una sola persona concentre todo el ciclo.",
    ])

    add_heading(doc, "5. Planeación y autorización del egreso", 1)
    add_stage(
        doc, 1, "Capturar el anteproyecto de egresos",
        "Presupuesto > Egreso > Planeación > Anteproyecto de Egresos\nRuta: /Presupuesto/Egreso/Planeacion/Anteproyecto_Egresos",
        "Planeación presupuestal",
        "Empresa y año presupuestal seleccionados; catálogos de programa, partida, área y clasificación activos.",
        "Registra la necesidad presupuestal por programa, partida, área y calendario mensual.",
        [
            "Abrir Anteproyecto de Egresos y elegir Nueva partida o el alta desde Excel, según el permiso disponible.",
            "Capturar Programa, Partida, Área, Fecha y Descripción.",
            "Completar Fuente de financiamiento, Tipo de gasto, Dígito identificador, Destino de gasto y PY cuando aplique.",
            "Distribuir el importe en los meses permitidos; revisar que el total corresponda a la planeación aprobada.",
            "Guardar y revisar el registro en Partidas del anteproyecto.",
            "Adjuntar el soporte documental cuando el procedimiento interno lo requiera.",
        ],
        [
            "Programa, partida y área son obligatorios.",
            "No registrar importes negativos; la suma mensual debe representar el total planeado.",
            "Antes de autorizar, confirmar que la clasificación presupuestal sea coherente y esté activa.",
            "Un anteproyecto autorizado queda protegido contra edición ordinaria.",
        ],
        "Partida de egreso proyectado lista para autorización.",
        "Mientras no esté autorizado, use Editar. Si ya fue autorizado, utilice Quitar autorización / Regresar a proyectado; el autorizado se conserva inactivo para trazabilidad.",
        "Programa, partida, área, fecha, descripción, FF, TG, DI, DG, PY e importes de enero a diciembre.",
        "La solicitud se materializa con la acción Autorizar de la fila. El sistema pide confirmación y abre Autorizar Presupuesto de Egresos; la acción final es Firmar y guardar.",
    )

    add_stage(
        doc, 2, "Autorizar el anteproyecto y generar presupuesto autorizado",
        "Desde Anteproyecto de Egresos, acción Autorizar; consulta posterior en /Presupuesto/Egreso/Presupuesto_Autorizado",
        "Autorizador de presupuesto",
        "Anteproyecto completo y revisado.",
        "Convierte el egreso proyectado en techo presupuestal autorizado y registra fecha/usuario de autorización.",
        [
            "Localizar la partida del anteproyecto y seleccionar el icono Autorizar.",
            "Confirmar el mensaje que indica que se generará el presupuesto autorizado.",
            "Revisar la información precargada, la distribución mensual y la clasificación.",
            "Seleccionar Firmar y guardar.",
            "Abrir Presupuesto Autorizado y comprobar que el registro aparezca vinculado al anteproyecto.",
            "Conservar la notificación y el soporte de autorización dentro del expediente.",
        ],
        [
            "No autorizar partidas incompletas o clasificadas en una fuente/área incorrecta.",
            "El sistema identifica los autorizados provenientes de anteproyecto.",
            "Para corregir no se elimina el antecedente: se usa Regresar a proyectado.",
        ],
        "Presupuesto autorizado disponible para las etapas de suficiencia y compromiso.",
        "Use Quitar autorización o Regresar a proyectado. Revise el efecto mostrado por el sistema antes de confirmar.",
        auth="La autorización crea un nuevo registro relacionado y bloquea la edición del anteproyecto. El usuario autorizador debe ser distinto del capturista cuando la política interna lo exija.",
    )

    add_heading(doc, "6. Adquisición: solicitud, mercado y suficiencia", 1)
    add_stage(
        doc, 3, "Crear la requisición",
        "Adquisiciones > Requisiciones\nRuta: /Adquisiciones/Requisicion",
        "Área solicitante / Compras",
        "Necesidad aprobada y clasificación presupuestal conocida.",
        "Formaliza la solicitud operativa de bienes o servicios y su vinculación presupuestal.",
        [
            "Crear una requisición y capturar descripción, área, solicitante, fecha, importe y tipo Bien/Servicio.",
            "Seleccionar proyecto y programa; completar FF, TG, DI y DG.",
            "Registrar oficio, fecha de oficio y observaciones cuando correspondan.",
            "Guardar, abrir Ver detalle y agregar al menos una partida presupuestal.",
            "Seleccionar la partida y agregar los bienes o servicios con cantidad, unidad y datos requeridos.",
            "Revisar que el importe de la cabecera sea mayor que cero y sea congruente con el detalle.",
        ],
        [
            "La clasificación completa es requisito para enviar a suficiencia.",
            "La requisición debe contener bienes/servicios y todos deben tener partida.",
            "Cuando ya existe una cotización activa, la requisición se muestra En cotización y queda bloqueada para editar o eliminar.",
            "No duplique la requisición para corregir una clasificación; libere o corrija el vínculo conforme al procedimiento autorizado.",
        ],
        "Requisición con partidas y bienes lista para cotizar.",
        "Corrija antes de crear cotizaciones. Si ya está bloqueada, libere/cancele el vínculo activo con la cotización conforme a permisos y política.",
        "Descripción, área, solicitante, fecha, importe, tipo, proyecto, programa, FF, TG, DI, DG, oficio y observaciones.",
    )

    add_stage(
        doc, 4, "Capturar y recibir cotizaciones",
        "Adquisiciones > Cotización\nRuta: /Adquisiciones/Cotizacion",
        "Compras / Adquisiciones",
        "Requisición con al menos un bien o servicio.",
        "Integra precios de proveedor y permite calcular el monto base de la solicitud de suficiencia.",
        [
            "Desde Requisiciones seleccionar Generar cotización, o crearla desde Cotización.",
            "Elegir requisición y proveedor; registrar fechas de solicitud, cotización y compromiso.",
            "Capturar entrega, vigencia, condiciones y comentarios; decidir si se envía la solicitud por correo.",
            "Abrir la cotización y usar Captura rápida de bienes para incorporar los renglones requeridos.",
            "Registrar Precio unitario para cada bien; usar Recibir montos cotizados para completar la respuesta del proveedor.",
            "Comparar vigencia, condiciones e importe; conservar las evidencias de cada proveedor.",
        ],
        [
            "Todos los bienes de la requisición deben tener por lo menos un monto cotizado antes de solicitar suficiencia.",
            "El precio unitario debe ser mayor que cero.",
            "No mezclar bienes de otra requisición ni proveedores inactivos.",
            "La solicitud de suficiencia usa el promedio de las cotizaciones capturadas más el porcentaje de ajuste indicado.",
        ],
        "Requisición totalmente cotizada y lista para solicitar suficiencia.",
        "Editar o eliminar el detalle cotizado mientras no haya avanzado el flujo; después documentar la corrección y regenerar la solicitud si procede.",
        "Requisición, proveedor, fechas, entrega, vigencia, condiciones, comentarios y precio unitario por bien.",
    )

    add_stage(
        doc, 5, "Generar la solicitud de suficiencia",
        "Desde Requisiciones, acción Enviar a suficiencia y autorización; seguimiento en /Adquisiciones/Solicitud_Suficiencia",
        "Compras / Área solicitante",
        "Requisición con clasificación completa, bienes con partida y todos los bienes cotizados.",
        "Solicita la validación y reserva presupuestal por partida y mes.",
        [
            "En Requisiciones seleccionar Enviar a suficiencia y autorización.",
            "Confirmar la requisición y la fecha de solicitud.",
            "Indicar el porcentaje de ajuste cuando esté autorizado; nunca usar un valor negativo.",
            "Revisar el precálculo: bienes, cantidad, número de cotizaciones, promedio y monto ajustado.",
            "Capturar la justificación y, si aplica, los datos de gasto no programable o compromiso de nómina.",
            "Generar la solicitud y verificar su detalle por bien, partida, mes y total.",
        ],
        [
            "No puede existir otra solicitud activa para la misma requisición, salvo que la anterior esté rechazada/inactiva según la regla del sistema.",
            "El sistema impide generar la solicitud si falta un bien, partida o monto cotizado.",
            "El mes aplicado se deriva de la fecha de la requisición.",
            "Revisar que el total solicitado sea congruente con el mercado y el presupuesto autorizado.",
        ],
        "Solicitud de suficiencia con detalle lista para autorización presupuestal.",
        "Mientras sea editable, ajuste la solicitud. Si fue rechazada, corrija la requisición/cotizaciones y genere o reactive el flujo conforme a la opción disponible.",
        "Requisición, fecha, porcentaje de ajuste, estatus, justificación y datos de gasto no programable.",
        "La acción crea la solicitud, pero no reserva presupuesto. La reserva ocurre únicamente al autorizarla en Presupuesto Comprometido.",
    )

    add_stage(
        doc, 6, "Autorizar la suficiencia presupuestal",
        "Presupuesto > Egreso > Presupuesto Comprometido > Solicitud Suficiencia\nRuta: /Presupuesto/Egreso/Presupuesto_Comprometido/Solicitud_Suficiencia",
        "Autorizador presupuestal",
        "Solicitud activa con detalle e importes mayores que cero; usuario relacionado con una persona autorizadora.",
        "Valida disponibilidad y crea la autorización con reserva presupuestal en una sola transacción.",
        [
            "Abrir Solicitud Suficiencia y localizar la solicitud pendiente.",
            "Abrir Ver detalle y revisar requisición, mes, bienes, partidas y total solicitado.",
            "Seleccionar Autorizar solicitud y confirmar el mensaje de generación de la autorización con detalle.",
            "Verificar Empresa, Solicitud suficiencia, Autorizado por, Fecha de autorización, Justificación y Observaciones.",
            "Confirmar la autorización. El sistema valida saldo por año, programa, partida, área, FF, TG, DI y DG.",
            "Comprobar que la solicitud cambie al estatus autorizado y que la autorización aparezca en Autorización Suficiencia.",
        ],
        [
            "No se autoriza una solicitud rechazada ni una solicitud sin detalle.",
            "No se permite más de una autorización activa para la misma solicitud.",
            "Si el importe solicitado supera el saldo neto disponible, el sistema devuelve Saldo insuficiente y no reserva.",
            "La persona autorizadora debe existir y estar activa.",
            "La autorización y sus detalles quedan bloqueados cuando alcanzan el estatus autorizado.",
        ],
        "Suficiencia autorizada y presupuesto reservado; requisición habilitada para orden de compra.",
        "No intente modificar la autorización bloqueada. Corrija el origen, cancele/rechace conforme al procedimiento y libere la reserva antes de emitir una nueva autorización.",
        auth="Esta es la autorización presupuestal obligatoria. El sistema copia el detalle de la solicitud y valida el saldo antes de confirmar la reserva.",
    )

    add_heading(doc, "7. Orden de compra y recepción", 1)
    add_stage(
        doc, 7, "Crear y autorizar la orden de compra",
        "Adquisiciones > Órdenes de Compra\nRuta: /Adquisiciones/Orden_Compra",
        "Compras / Autorizador de adquisiciones",
        "Requisición activa con suficiencia autorizada; proveedor seleccionado.",
        "Emite la instrucción de compra por proveedor, bienes, partidas e importes.",
        [
            "Crear la orden y seleccionar requisición y proveedor.",
            "Capturar número, descripción, fecha de OC, fechas requerida/entrega/vigencia, moneda, tipo de cambio y compra directa cuando aplique.",
            "Abrir Ver detalle y agregar los renglones desde bienes cotizados o bienes de requisición.",
            "Para cada renglón registrar cantidad, unidad, precio unitario, IVA y observaciones.",
            "Agregar las partidas presupuestales con fuente e importe.",
            "Verificar que exista al menos un detalle y una partida, y que el total de partidas coincida con el total de detalles.",
            "Seleccionar Autorizar y confirmar. La orden cambia de INICIAL a POR SURTIR y queda bloqueada para edición/eliminación.",
        ],
        [
            "La fecha de OC debe ser igual o posterior a la fecha de requisición.",
            "Cantidad y precio unitario deben ser mayores que cero.",
            "El proveedor debe estar activo y pertenecer a la empresa del proceso.",
            "La requisición debe tener una solicitud autorizada y una autorización activa.",
            "No autorizar si el total de partidas no coincide con el total de los renglones.",
        ],
        "Orden autorizada en estatus POR SURTIR, disponible para recepción y registro de compromiso.",
        "Antes de autorizar, edite detalles/partidas. Después de autorizar, la orden queda bloqueada; cualquier cambio requiere el procedimiento formal de cancelación o reversa.",
        "Número, estatus, descripción, requisición, proveedor, fechas, moneda, tipo de cambio, compra directa, renglones y partidas.",
        "La autorización de la orden es distinta de la suficiencia: la primera permite surtir; la segunda reservó el presupuesto.",
    )

    add_stage(
        doc, 8, "Registrar recepción y entrada de bienes",
        "Almacén > Recepción de pedidos\nRuta: /Almacen/Recepcion_Pedidos",
        "Almacén / Recepción",
        "Orden de compra autorizada y bienes entregados por el proveedor.",
        "Registra entradas trazables contra cada renglón de la orden y controla cantidades recibidas/pendientes.",
        [
            "Abrir Recepción de pedidos y seleccionar la orden de compra.",
            "En Renglones por recibir revisar bien, cantidad solicitada, recibida, pendiente y precio.",
            "Seleccionar Registrar entrada en el renglón correspondiente.",
            "Confirmar Bien o servicio, Área/ubicación, Cantidad, Costo unitario, Fecha de entrada y Unidad.",
            "Capturar Motivo, Caducidad, Clave, Factura, Remisión y Lote cuando correspondan.",
            "Guardar y abrir Ver entradas para comprobar la clave, cantidad, costo, fecha y documento.",
            "Repetir para entregas parciales hasta que la cantidad pendiente llegue a cero.",
        ],
        [
            "La cantidad de entrada debe ser mayor que cero y no debe rebasar la pendiente de la orden.",
            "El costo sugerido proviene del precio unitario; valide el documento del proveedor antes de guardar.",
            "No mezclar lotes, remisiones o facturas cuando deban conservarse por separado.",
            "El botón Registrar entrada se deshabilita cuando no hay cantidad pendiente.",
            "La entrada es un movimiento de almacén; la alta patrimonial puede requerir un paso adicional para bienes capitalizables.",
        ],
        "Movimiento de almacén ligado al detalle de la orden; cantidades recibidas y pendientes actualizadas.",
        "No sobrescriba una entrada cerrada/contabilizada. Use el ajuste contrario autorizado y conserve la referencia al movimiento original.",
        "Bien/servicio, ubicación, cantidad, costo, fecha, unidad, motivo, caducidad, clave, factura, remisión y lote.",
    )

    add_callout(doc, "Bienes patrimoniales", "Si el bien debe controlarse como activo fijo, complete el alta en Patrimonio > Bienes conforme a la política de capitalización. El sistema dispone de una relación con el detalle de la orden; valide clave, factura, valor, ubicación y resguardo.", "info")

    add_heading(doc, "8. Compromiso, factura, CLC y pago", 1)
    add_body(doc, "Después de la suficiencia y la orden autorizada, el expediente continúa por la ruta financiera. Dependiendo de la configuración institucional puede usarse el Registro Comprometido derivado de la autorización de suficiencia y/o el Registro de Compromiso contractual vinculado a la orden autorizada. No duplique compromisos para el mismo objeto.")

    add_stage(
        doc, 9, "Registrar compromiso o contrato",
        "Presupuesto > Egreso > Presupuesto Comprometido > Registro Comprometido\nRuta: /Presupuesto/Egreso/Presupuesto_Comprometido/Registro_Comprometido\nAlterna contractual: /Adquisiciones/Contratos/Registro_Compromiso",
        "Presupuesto / Contratos",
        "Autorización de suficiencia activa y, para el registro ORCO, orden de compra autorizada.",
        "Formaliza el compromiso presupuestal/contractual y prepara la recepción de factura.",
        [
            "Desde Autorización Suficiencia revisar el detalle autorizado y usar Generar registro comprometido cuando corresponda.",
            "Confirmar número, descripción, fechas, monto total y calendario por partida.",
            "Si se utiliza Registro de Compromiso de Adquisiciones, seleccionar la orden autorizada, proveedor, vigencia y montos mínimo/máximo.",
            "Adjuntar contrato, convenio, orden autorizada y soporte de suficiencia.",
            "Autorizar el registro contractual cuando esté completo; el sistema lo bloquea para edición operativa.",
            "Verificar que el monto del compromiso no exceda la suficiencia ni la orden.",
        ],
        [
            "La autorización de suficiencia debe estar autorizada y tener detalle.",
            "La orden debe estar autorizada antes del compromiso contractual.",
            "No generar un segundo contrato activo para la misma autorización cuando el sistema ya reporta uno.",
            "La fecha final de vigencia no puede ser anterior a la inicial.",
        ],
        "Contrato/compromiso activo con calendario por partida, listo para factura.",
        "Corrija mientras esté en borrador/inicial. Una vez autorizado, aplique convenio, cancelación o proceso formal; no edite el registro bloqueado.",
        auth="La autorización contractual bloquea el registro. La autorización presupuestal previa continúa siendo la fuente de la reserva.",
    )

    add_stage(
        doc, 10, "Registrar factura y enviarla a autorización",
        "Tesorería > Cuentas por pagar > Recepción de Facturas y Comprobantes de Pago\nRuta: /Presupuesto/Tesoreria/CuentasXPagar/Factura_Pago",
        "Cuentas por pagar",
        "Contrato/compromiso, póliza y factura válida del proveedor.",
        "Registra el documento fiscal y sus partidas para generar la CLC.",
        [
            "Crear o generar la factura desde Registro Comprometido mediante Recepción de factura.",
            "Seleccionar Empresa, Contrato y Póliza.",
            "Capturar número y serie de factura, fechas de emisión/recepción, subtotal, IVA, retención, total y UUID.",
            "Agregar o revisar las partidas de factura y el monto aplicado a cada una.",
            "Validar que subtotal, impuestos, retenciones y total coincidan con el CFDI/documento recibido.",
            "Seleccionar Enviar a autorización. El sistema prepara la CLC y sus detalles a partir de la factura.",
        ],
        [
            "La factura debe tener detalle antes de generar CLC.",
            "Contrato y póliza son obligatorios en la pantalla actual.",
            "No aplicar una misma partida por encima del monto facturado/comprometido.",
            "Después de enviar a autorización, la factura queda bloqueada para edición ordinaria.",
        ],
        "Factura vinculada a contrato y póliza; CLC preparada para autorización/provisión.",
        "Antes de enviar, edite o elimine la partida incorrecta. Después, corrija mediante cancelación/reversa de CLC y documente el motivo.",
        "Contrato, póliza, número/serie, fechas, subtotal, IVA, retención, total, UUID, estatus y observaciones.",
        auth="La acción Enviar a autorización genera la CLC con el detalle presupuestal y las facturas asociadas; no equivale todavía a la salida bancaria.",
    )

    add_stage(
        doc, 11, "Revisar CLC y provisionar el pago",
        "Tesorería > Cuentas por pagar > Provisión del Pago\nRuta: /Presupuesto/Tesoreria/CuentasXPagar/Provision_Pago",
        "Cuentas por pagar / Tesorería",
        "CLC generada desde factura con contrato, póliza y detalle presupuestal.",
        "Concentra el calendario, partidas y facturas autorizadas que serán pagadas.",
        [
            "Abrir Provisión del Pago y localizar la CLC.",
            "Revisar número CLC, contrato, proveedor, fecha de solicitud, importe y estatus.",
            "Abrir el detalle y validar partidas, total por partida y facturas asociadas.",
            "Confirmar fechas de solicitud/autorización, póliza e importe total.",
            "Seleccionar Elaborar cheque o transferencia (icono de pagos).",
            "Revisar la información precargada del cheque y sus partidas antes de guardar.",
        ],
        [
            "La CLC debe contener por lo menos un detalle.",
            "El total de la CLC debe coincidir con la suma de sus detalles y facturas aplicadas.",
            "No pagar una CLC cancelada, rechazada o ya provisionada.",
            "Al generar el cheque/transferencia, la CLC avanza a un estatus bloqueado para edición.",
        ],
        "CLC provisionada y cheque/transferencia en preparación.",
        "Si detecta diferencia antes de generar el pago, corrija detalle/factura. Después de generarlo, cancele o reverse conforme al proceso de Tesorería.",
        "Empresa, contrato, póliza, número CLC, fechas, importe total, estatus, observaciones, detalle y facturas.",
    )

    add_stage(
        doc, 12, "Emitir cheque o transferencia",
        "Tesorería > Cuentas por pagar > Elaboración de Cheques o Transferencias\nRuta: /Presupuesto/Tesoreria/CuentasXPagar/Cheque_Transferencia",
        "Tesorería",
        "CLC provisionada, cuenta bancaria disponible, póliza y autorización de pago.",
        "Registra el medio de pago, su importe y las partidas pagadas.",
        [
            "Desde la CLC abrir Provisión de pago o crear el registro en Elaboración de Cheques o Transferencias.",
            "Seleccionar Empresa, CLC, Cuenta bancaria y Póliza.",
            "Capturar Fecha de emisión, número de Cheque/transferencia, Concepto, Importe total y Observaciones.",
            "Revisar las partidas precargadas y sus montos pagados.",
            "Guardar el registro y comprobar número, CLC, contrato, fecha, importe y estatus en el listado.",
            "Adjuntar comprobante bancario, póliza y evidencia de entrega/aplicación.",
        ],
        [
            "La cuenta bancaria y la póliza son obligatorias.",
            "El importe del pago debe coincidir con la suma de partidas y no exceder el saldo autorizado de la CLC.",
            "Una provisión/pago autorizado queda bloqueado para modificar o eliminar.",
            "Conciliar número de cheque/transferencia, fecha y movimiento bancario.",
        ],
        "Cheque o transferencia registrado con partidas pagadas y expediente de soporte.",
        "Antes de autorizar/conciliar, edite. Después, use cancelación o reversa autorizada y conserve la relación con el pago original.",
        "CLC, cuenta bancaria, póliza, fecha, número, concepto, importe, estatus, observaciones y partidas.",
        auth="La salida bancaria sólo debe ejecutarse por un usuario autorizado y con evidencia de revisión de la CLC, la cuenta y el beneficiario.",
    )

    add_heading(doc, "9. Matriz de autorización y bloqueo", 1)
    add_table(doc,
              ["Registro", "Solicita / prepara", "Autoriza", "Efecto", "Cómo corregir"],
              [
                  ["Anteproyecto", "Planeación", "Autorizador de presupuesto", "Crea presupuesto autorizado; protege el proyectado.", "Regresar a proyectado."],
                  ["Solicitud suficiencia", "Compras / solicitante", "Autorizador presupuestal", "Reserva saldo y copia el detalle autorizado.", "Rechazo/cancelación y nueva solicitud."],
                  ["Orden de compra", "Compras", "Autorizador de adquisiciones", "Pasa de INICIAL a POR SURTIR; bloquea edición.", "Cancelación/reversa formal."],
                  ["Contrato/compromiso", "Contratos / Presupuesto", "Autorizador contractual", "Bloquea el compromiso operativo.", "Convenio, cancelación o reversa."],
                  ["Factura a CLC", "Cuentas por pagar", "Revisor/autorizador de CLC", "Genera CLC y bloquea factura según estatus.", "Cancelar/revertir CLC."],
                  ["CLC a pago", "Cuentas por pagar", "Tesorería", "Genera cheque/transferencia y bloquea CLC.", "Cancelar/revertir pago."],
              ],
              [1550, 1600, 1700, 2580, 1930], font_size=8.1)

    add_heading(doc, "10. Conciliación de importes y documentos", 1)
    add_table(doc,
              ["Comparación", "Debe cumplirse", "Evidencia"],
              [
                  ["Anteproyecto vs. autorizado", "Misma clasificación y calendario salvo ajuste autorizado.", "Detalle mensual y registro de autorización."],
                  ["Requisición vs. cotización", "Todos los bienes cotizados; cantidades y unidades comparables.", "Cotizaciones y respuestas de proveedor."],
                  ["Solicitud vs. autorización", "Mismo detalle; total autorizado no mayor al disponible.", "Solicitud, autorización y saldo."],
                  ["Orden: detalles vs. partidas", "La suma de partidas coincide con el total de detalles.", "Orden autorizada."],
                  ["Orden vs. entradas", "Recibido acumulado no mayor a solicitado; pendiente correcto.", "Remisión, factura, lote y movimientos de almacén."],
                  ["Contrato vs. factura", "Factura dentro de vigencia y saldo comprometido.", "Contrato, CFDI/UUID y póliza."],
                  ["Factura/CLC/cheque", "Totales y partidas coinciden; pago no excede CLC.", "CLC, cheque/transferencia y comprobante bancario."],
              ],
              [2200, 4200, 2960], font_size=8.7)

    add_heading(doc, "11. Expediente documental mínimo", 1)
    add_checklist(doc, [
        "Justificación y anteproyecto de egresos.",
        "Evidencia de presupuesto autorizado.",
        "Requisición con partidas y detalle de bienes/servicios.",
        "Cotizaciones, comparativo y selección de proveedor.",
        "Solicitud y autorización de suficiencia.",
        "Orden de compra autorizada.",
        "Contrato/compromiso y anexos, cuando aplique.",
        "Remisión, recepción y entradas de almacén.",
        "Factura/CFDI, UUID y póliza.",
        "CLC autorizada/provisionada.",
        "Cheque o transferencia y comprobante bancario.",
        "Acta, resguardo o alta patrimonial para bienes capitalizables.",
    ])

    add_heading(doc, "12. Correcciones, rechazos y cancelaciones", 1)
    add_bullets(doc, [
        "Corregir en el registro origen mientras siga editable; no crear duplicados para ocultar un error.",
        "Si existe un vínculo activo (cotización, suficiencia, orden, CLC o pago), revertir primero el documento posterior conforme a permisos.",
        "Registrar motivo, usuario, fecha y documento afectado en observaciones y soporte documental.",
        "Cuando el sistema conserva un registro inactivo para trazabilidad, no eliminarlo de la evidencia del expediente.",
        "Las entradas de almacén cerradas o contabilizadas se corrigen con un movimiento contrario, no sobrescribiendo el original.",
        "Después de una cancelación, verificar que el saldo presupuestal, pendiente de orden, saldo de contrato y saldo bancario queden consistentes.",
    ])
    add_callout(doc, "Escalamiento", "Si la pantalla está bloqueada y no existe una acción de reversa permitida, detenga el flujo y solicite intervención al responsable funcional. No modifique datos directamente en base de datos como sustituto del proceso.", "risk")

    add_heading(doc, "13. Lista de cierre del proceso", 1)
    add_checklist(doc, [
        "El presupuesto autorizado conserva vínculo con el anteproyecto.",
        "La requisición tiene clasificación completa y detalle correcto.",
        "Todos los bienes cuentan con cotización y soporte de proveedor.",
        "La suficiencia está autorizada y el saldo fue reservado.",
        "La orden está autorizada y sus detalles cuadran con partidas.",
        "Las entradas de almacén no exceden las cantidades solicitadas.",
        "El compromiso/contrato corresponde a la orden y suficiencia.",
        "La factura y sus partidas coinciden con lo recibido y contratado.",
        "La CLC incluye las facturas y partidas correctas.",
        "El cheque/transferencia coincide con la CLC y la cuenta bancaria.",
        "El expediente documental está completo y consultable.",
        "Los bienes capitalizables quedaron dados de alta y, si aplica, resguardados.",
    ])

    add_heading(doc, "14. Glosario breve", 1)
    add_table(doc,
              ["Término", "Definición operativa"],
              [
                  ["Anteproyecto de egresos", "Planeación mensual del gasto por programa, partida, área y clasificación."],
                  ["Suficiencia", "Validación de que existe saldo disponible para atender una requisición."],
                  ["Reserva", "Afectación que aparta presupuesto para una autorización de suficiencia."],
                  ["Compromiso", "Obligación presupuestal/contractual derivada de una autorización y una compra."],
                  ["Orden de compra", "Instrucción autorizada al proveedor con bienes, cantidades, precios y partidas."],
                  ["Entrada de almacén", "Movimiento que reconoce la recepción física contra un renglón de orden."],
                  ["CLC", "Cuenta por Liquidar Certificada que concentra factura, partidas, póliza e importe a pagar."],
                  ["Cheque/transferencia", "Registro del medio de pago vinculado a una CLC y sus partidas."],
                  ["FF / TG / DI / DG", "Fuente de financiamiento / Tipo de gasto / Dígito identificador / Destino del gasto."],
              ],
              [2200, 7160], font_size=9.1)

    add_callout(doc, "Fin del manual", "Use este documento como guía operativa y compleméntelo con las políticas de autorización, montos, firmas y archivo vigentes en la organización.", "info")

    # Core properties and update-fields flag.
    doc.core_properties.title = "Manual del flujo de anteproyecto de egresos a cheque y entrada de bienes"
    doc.core_properties.subject = "Solicitud, autorización, adquisición, recepción y pago en eGestion360"
    doc.core_properties.author = "Gestión Empresarial 360"
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_document()
